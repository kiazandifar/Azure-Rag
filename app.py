import os
import time
import logging
import tempfile
import numpy as np
import pandas as pd
import tiktoken
import re
from flask import Flask, render_template, request, redirect, url_for, session, flash
from openai import AzureOpenAI
from dotenv import load_dotenv
from docx import Document
import pdfplumber

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Replace with a secure secret key

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Initialize the OpenAI client
client = AzureOpenAI(
    api_key=os.getenv("AZURE_API_KEY"),
    api_version=os.getenv("AZURE_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_ENDPOINT")
)

# Initialize the tokenizer
tokenizer = tiktoken.get_encoding("cl100k_base")

# Function to extract text and tables from DOCX files
def extract_text_and_tables_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        # Convert table to string
        df = pd.DataFrame(table_data)
        table_string = df.to_csv(sep='|', index=False, header=False)
        full_text.append(table_string)
    return '\n'.join(full_text)

# Function to extract text and tables from PDF files
def extract_text_and_tables_from_pdf(file_path):
    full_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract text
            text = page.extract_text()
            if text:
                full_text.append(text)
            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                df = pd.DataFrame(table)
                table_string = df.to_csv(sep='|', index=False, header=False)
                full_text.append(table_string)
    return '\n'.join(full_text)

# Function to normalize text
def normalize_text(s):
    s = re.sub(r'\s+', ' ', s).strip()
    s = re.sub(r"\. ,", "", s)
    s = s.replace("..", ".")
    s = s.replace(". .", ".")
    s = s.strip()
    return s

# Function to split text into chunks based on token limit
def split_text_into_chunks(text, max_tokens=8000, overlap=200):
    tokens = tokenizer.encode(text)
    chunks = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = tokenizer.decode(chunk_tokens)
        chunks.append(chunk_text)
        start += max_tokens - overlap  # Move to the next chunk with overlap
    return chunks

@app.route('/', methods=['GET'])
def index():
    # Initialize conversation if not present
    if 'conversation' not in session:
        session['conversation'] = []
        logging.debug("Initialized conversation in session.")
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        flash('No files uploaded.')
        logging.warning("No files uploaded.")
        return redirect(url_for('index'))

    files = request.files.getlist('files')

    if not files:
        flash('No files selected.')
        logging.warning("No files selected.")
        return redirect(url_for('index'))

    # Save uploaded files to a temporary directory
    upload_folder = 'uploads'
    os.makedirs(upload_folder, exist_ok=True)
    file_paths = []
    for file in files:
        filename = file.filename
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)
        file_paths.append(file_path)
        logging.debug(f"Saved uploaded file: {file_path}")

    # Process files to extract text including tables
    documents = []
    for path in file_paths:
        if path.lower().endswith(('.doc', '.docx')):
            text = extract_text_and_tables_from_docx(path)
        elif path.lower().endswith('.pdf'):
            text = extract_text_and_tables_from_pdf(path)
        else:
            # For other file types, read as plain text
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        text = normalize_text(text)
        # Split text into chunks
        chunks = split_text_into_chunks(text)
        for idx, chunk in enumerate(chunks):
            temp_file = tempfile.NamedTemporaryFile(
                delete=False,
                mode='w',
                encoding='utf-8',
                suffix=f'_{idx}.txt',
                dir=upload_folder
            )
            temp_file.write(chunk)
            temp_file.close()
            documents.append(temp_file.name)
            logging.debug(f"Created temporary chunk file: {temp_file.name}")

    # Create a vector store with the name of the first file + "_vector_store"
    vector_store_name = os.path.splitext(files[0].filename)[0] + "_vector_store"
    vector_store = client.beta.vector_stores.create(name=vector_store_name)
    logging.debug(f"Created vector store with ID: {vector_store.id}")

    # Store the vector store ID in the session
    session['vector_store_id'] = vector_store.id

    # Open file streams
    file_streams = [open(path, "rb") for path in documents]

    # Upload documents to the vector store and poll for completion
    file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
        vector_store_id=vector_store.id, files=file_streams
    )
    logging.debug("Uploaded documents to vector store and completed processing.")

    # Close file streams
    for file_stream in file_streams:
        file_stream.close()

    # Clean up temporary files
    for path in documents:
        os.remove(path)
        logging.debug(f"Removed temporary file: {path}")

    # Clean up uploaded files
    for path in file_paths:
        os.remove(path)
        logging.debug(f"Removed uploaded file: {path}")

    flash('Documents uploaded and vector store created successfully.')
    return redirect(url_for('index'))

@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.form['user_input']
    logging.debug(f"Received user input: {user_input}")
    vector_store_id = session.get('vector_store_id')

    if not vector_store_id:
        flash('Vector store not initialized. Please upload documents first.')
        logging.warning("Vector store ID not found in session.")
        return redirect(url_for('index'))

    # Retrieve assistant ID from the session
    assistant_id = session.get('assistant_id')

    if not assistant_id:
        # Initialize the assistant with the vector store
        assistant = client.beta.assistants.create(
            instructions="You are a helpful assistant that provides information based on uploaded documents, including data from tables within the documents.",
            model="gpt-4o-appl",  # Replace with your model deployment name
            temperature=0,
            tools=[{"type": "file_search"}],
            tool_resources={"file_search": {"vector_store_ids": [vector_store_id]}}
        )
        assistant_id = assistant.id
        session['assistant_id'] = assistant_id
        logging.debug(f"Created new assistant with ID: {assistant_id}")
    else:
        logging.debug(f"Using existing assistant ID: {assistant_id}")

    # Create a new thread for conversation
    thread = client.beta.threads.create()
    thread_id = thread.id
    session['thread_id'] = thread_id
    logging.debug(f"Created new thread with ID: {thread_id}")

    # Add user message to the thread
    client.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=user_input
    )
    logging.debug("Added user message to thread.")

    # Run the assistant
    run = client.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id
    )
    logging.debug(f"Started assistant run with ID: {run.id}")

    # Polling until the run completes
    while run.status in ['queued', 'in_progress', 'cancelling']:
        logging.debug(f"Run status: {run.status}. Waiting...")
        time.sleep(1)
        run = client.beta.threads.runs.retrieve(
            thread_id=thread_id,
            run_id=run.id
        )

    if run.status == 'completed':
        messages = client.beta.threads.messages.list(
            thread_id=thread_id
        )
        # Convert messages to a list
        messages_list = list(messages)
        logging.debug(f"Retrieved {len(messages_list)} messages from thread.")
        for msg in messages_list:
            logging.debug(f"Message ID: {msg.id}, Role: {msg.role}, Content: {msg.content}")

        # Extract the latest assistant message
        response_text = ""
        for msg in reversed(messages_list):
            if msg.role == 'assistant':
                logging.debug(f"Found assistant message with content: {msg.content}")
                for block in msg.content:
                    response_text += block.text.value + "\n"
                break

        # Update conversation in session
        session['conversation'].append({'role': 'user', 'content': user_input})
        session['conversation'].append({'role': 'assistant', 'content': response_text.strip()})
        # Limit conversation history to last 20 entries (10 exchanges)
        session['conversation'] = session['conversation'][-20:]
        session.modified = True  # Mark session as modified
        logging.debug("Updated conversation in session.")
    else:
        flash(f"An error occurred: {run.status}")
        logging.error(f"Assistant run failed with status: {run.status}")

    return redirect(url_for('index'))

@app.route('/delete_vector_store', methods=['POST'])
def delete_vector_store():
    vector_store_id = session.get('vector_store_id')
    assistant_id = session.get('assistant_id')
    thread_id = session.get('thread_id')

    if not vector_store_id:
        flash('No vector store to delete.')
        logging.warning("No vector store ID found in session.")
        return redirect(url_for('index'))

    try:
        # Delete the vector store
        client.beta.vector_stores.delete(vector_store_id=vector_store_id)
        session.pop('vector_store_id', None)
        logging.debug(f"Deleted vector store with ID: {vector_store_id}")

        # Delete the assistant if it exists
        if assistant_id:
            client.beta.assistants.delete(assistant_id=assistant_id)
            session.pop('assistant_id', None)
            logging.debug(f"Deleted assistant with ID: {assistant_id}")

        # Delete the thread if it exists
        if thread_id:
            client.beta.threads.delete(thread_id=thread_id)
            session.pop('thread_id', None)
            logging.debug(f"Deleted thread with ID: {thread_id}")

        # Clear conversation history
        session['conversation'] = []
        logging.debug("Cleared conversation history.")

        flash('Vector store, assistant, and thread deleted successfully.')
    except Exception as e:
        flash(f'Error deleting resources: {str(e)}')
        logging.error(f"Error deleting resources: {str(e)}")

    return redirect(url_for('index'))

@app.route('/reset', methods=['POST'])
def reset():
    # Clear the conversation history and assistant data
    session['conversation'] = []
    session.pop('assistant_id', None)
    session.pop('thread_id', None)
    flash('Conversation reset.')
    logging.debug("Conversation reset by user.")
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True)
